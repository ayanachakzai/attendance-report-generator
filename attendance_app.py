"""
ATTENDANCE REPORT GENERATOR - WEB APP (Full Version)
====================================================
Works on Streamlit Cloud with PDF support via LibreOffice
"""

import streamlit as st
import pandas as pd
from docx import Document
import os
import re
import subprocess
import tempfile
import zipfile
from io import BytesIO
import shutil

st.set_page_config(
    page_title="Attendance Report Generator",
    page_icon="📊",
    layout="wide"
)

def check_libreoffice():
    """Check if LibreOffice is installed"""
    # Check common LibreOffice paths
    possible_paths = [
        '/usr/bin/soffice',  # Linux/Streamlit Cloud
        '/Applications/LibreOffice.app/Contents/MacOS/soffice',  # Mac
        'C:\\Program Files\\LibreOffice\\program\\soffice.exe'  # Windows
    ]
    
    for path in possible_paths:
        if os.path.exists(path):
            return path
    
    # Try finding it in PATH
    if shutil.which('soffice'):
        return shutil.which('soffice')
    
    return None


def clean_text(value):
    """Return a trimmed string while treating blank spreadsheet cells as empty."""
    if pd.isna(value):
        return ""
    return str(value).replace("\xa0", " ").strip()


def normalize_label(value):
    value = clean_text(value).lower()
    value = re.sub(r"[\s_./\\-]+", " ", value)
    return re.sub(r"[^a-z0-9% ]+", "", value).strip()


def compact_label(value):
    return re.sub(r"[^a-z0-9%]+", "", normalize_label(value))


def is_bnu_header(value):
    compact = compact_label(value)
    return compact in {"bnuid", "bnuno", "bnunumber", "bnumber"} or (
        "bnu" in compact and any(part in compact for part in ["id", "no", "number"])
    )


def is_name_header(value):
    label = normalize_label(value)
    compact = compact_label(value)
    if is_surname_header(value):
        return False
    return compact in {"name", "firstname", "forename", "givenname", "preferredname"} or label in {
        "first name",
        "given name",
        "student first name",
    }


def is_surname_header(value):
    label = normalize_label(value)
    compact = compact_label(value)
    return compact in {"surname", "lastname", "familyname"} or label in {"last name", "family name"}


def is_full_name_header(value):
    label = normalize_label(value)
    compact = compact_label(value)
    return compact in {"fullname", "studentname", "learnername", "displayname"} or label in {
        "full name",
        "student name",
        "learner name",
        "name of student",
    }


def is_attendance_header(value):
    label = normalize_label(value)
    compact = compact_label(value)
    return (
        compact
        in {
            "live",
            "att",
            "att%",
            "attpercent",
            "attendance",
            "attendance%",
            "attendancepercent",
            "attendancepercentage",
            "attendancerate",
            "overall",
            "overall%",
            "overallattendance",
            "overallatt",
            "overallatt%",
            "actual",
            "actual%",
            "actualattendance",
            "present",
            "present%",
            "presentpercent",
            "presentpercentage",
            "percentage",
            "rate",
            "totalattendance",
        }
        or "attendance" in label
        or "attend" in label
        or "percent" in label
        or "percentage" in label
        or "present" in label
        or "overall" in label
        or "%" in label
    )


def is_campus_header(value):
    compact = compact_label(value)
    return compact in {"campus", "site", "location", "centre", "center"}


def is_group_header(value):
    label = normalize_label(value)
    compact = compact_label(value)
    return compact in {"group", "groupref", "class", "cohort", "intake"} or label in {
        "group ref",
        "group reference",
        "class group",
    }


def make_unique_columns(headers):
    columns = []
    seen = {}

    for index, value in enumerate(headers, start=1):
        name = clean_text(value) or f"Column {index}"
        count = seen.get(name, 0)
        seen[name] = count + 1
        columns.append(name if count == 0 else f"{name} {count + 1}")

    return columns


def parse_attendance_percent(value):
    if pd.isna(value):
        return None

    if isinstance(value, str):
        text = value.strip()
        if not text:
            return None
        lowered = text.lower()
        category_values = {
            "excellent": 80,
            "very good": 70,
            "good": 60,
            "could be better": 0,
            "poor": 0,
        }
        for phrase, percent in category_values.items():
            if phrase in lowered:
                return percent

        match = re.search(r"-?\d+(?:\.\d+)?", text.replace(",", ""))
        if not match:
            return None
        number = float(match.group())
        return number if "%" in text or number > 1 else number * 100

    try:
        number = float(value)
    except (TypeError, ValueError):
        return None

    return number if number > 1 else number * 100


def attendance_category_from_percent(attendance_percent):
    if attendance_percent >= 80:
        return "Excellent attendance"
    if attendance_percent >= 70:
        return "Very good attendance"
    if attendance_percent >= 60:
        return "Good attendance"
    return "Attendance could be better"


def clean_bnu_id(value):
    text = clean_text(value)
    if not text:
        return ""

    try:
        number = float(text)
        if number.is_integer() and re.fullmatch(r"\d+(?:\.0+)?", text):
            return str(int(number))
    except ValueError:
        pass

    return text.upper()


def choose_column(df, matcher, data_score=None):
    best_column = None
    best_score = 0

    for column in df.columns:
        score = 4 if matcher(column) else 0
        if data_score is not None:
            score += data_score(df[column])
        if score > best_score:
            best_column = column
            best_score = score

    return best_column if best_score >= 4 else None


def attendance_data_score(series):
    values = [parse_attendance_percent(value) for value in series.dropna().head(50)]
    valid_values = [value for value in values if value is not None and 0 <= value <= 100]
    return min(len(valid_values), 5)


def choose_attendance_column(df):
    best_column = None
    best_score = 0

    for column in df.columns:
        if any(matcher(column) for matcher in [is_bnu_header, is_name_header, is_surname_header, is_campus_header, is_group_header]):
            continue

        header_score = 5 if is_attendance_header(column) else 0
        data_score = attendance_data_score(df[column])
        score = header_score + data_score

        if score > best_score:
            best_column = column
            best_score = score

    return best_column if best_score >= 5 else None


def header_row_score(row):
    values = list(row)
    score = 0
    score += 5 if any(is_bnu_header(value) for value in values) else 0
    score += 3 if any(is_attendance_header(value) for value in values) else 0
    score += 2 if any(is_name_header(value) for value in values) else 0
    score += 2 if any(is_surname_header(value) for value in values) else 0
    score += 2 if any(is_full_name_header(value) for value in values) else 0
    score += 1 if any(is_campus_header(value) for value in values) else 0
    score += 1 if any(is_group_header(value) for value in values) else 0
    return score


def find_best_sheet_and_header(raw_sheets):
    best = None

    for sheet_name, raw in raw_sheets.items():
        max_rows = min(len(raw), 30)
        for row_number in range(max_rows):
            score = header_row_score(raw.iloc[row_number])
            if best is None or score > best["score"]:
                best = {"sheet_name": sheet_name, "row_number": row_number, "score": score}

    if not best or best["score"] < 8:
        raise ValueError(
            "Could not find the header row. The sheet must include BNU ID, name, and attendance columns."
        )

    return best["sheet_name"], best["row_number"]


def read_attendance_data(excel_file):
    if hasattr(excel_file, "seek"):
        excel_file.seek(0)

    raw_sheets = pd.read_excel(excel_file, sheet_name=None, header=None, dtype=object)
    sheet_name, header_row = find_best_sheet_and_header(raw_sheets)
    raw = raw_sheets[sheet_name]

    df = raw.iloc[header_row + 1:].copy()
    df.columns = make_unique_columns(raw.iloc[header_row])
    df = df.dropna(axis=1, how="all")
    df = df.dropna(how="all")

    bnu_column = choose_column(df, is_bnu_header)
    attendance_column = choose_attendance_column(df)
    name_column = choose_column(df, is_name_header)
    surname_column = choose_column(df, is_surname_header)
    full_name_column = choose_column(df, is_full_name_header)
    campus_column = choose_column(df, is_campus_header)
    group_column = choose_column(df, is_group_header)

    if bnu_column is None:
        raise ValueError("Could not find a BNU ID column. Column names are matched case-insensitively.")
    if attendance_column is None:
        column_list = ", ".join(clean_text(column) for column in df.columns)
        raise ValueError(
            "Could not find an attendance column. "
            f"I detected sheet '{sheet_name}', header row {header_row + 1}, and these columns: {column_list}. "
            "Please make sure the attendance percentage column has values like 0.82, 82, or 82%."
        )
    if name_column is None and full_name_column is None:
        raise ValueError("Could not find a student name column.")

    prepared = pd.DataFrame()
    prepared["BNU ID"] = df[bnu_column].apply(clean_bnu_id)

    if name_column is not None:
        prepared["Name"] = df[name_column].apply(clean_text)
    else:
        prepared["Name"] = df[full_name_column].apply(clean_text)

    prepared["Surname"] = df[surname_column].apply(clean_text) if surname_column is not None else ""

    if full_name_column is not None and name_column is None:
        prepared["Student Name"] = df[full_name_column].apply(clean_text)
    else:
        prepared["Student Name"] = (
            prepared["Name"].astype(str).str.strip() + " " + prepared["Surname"].astype(str).str.strip()
        ).str.strip()

    prepared["Campus"] = df[campus_column].apply(clean_text) if campus_column is not None else ""
    prepared["Group Ref"] = df[group_column].apply(clean_text) if group_column is not None else "Ungrouped"
    prepared["Attendance %"] = df[attendance_column].apply(parse_attendance_percent)
    prepared["LIVE"] = prepared["Attendance %"] / 100

    prepared = prepared[prepared["BNU ID"] != ""]
    prepared = prepared.dropna(subset=["Attendance %"])
    prepared = prepared[prepared["Student Name"] != ""]

    if prepared.empty:
        raise ValueError("No valid student rows were found after reading the attendance sheet.")

    prepared["Attendance Category"] = prepared["Attendance %"].apply(attendance_category_from_percent)
    prepared["Group Ref"] = prepared["Group Ref"].replace("", "Ungrouped")
    prepared.attrs["source_sheet"] = sheet_name
    prepared.attrs["header_row"] = header_row + 1
    prepared.attrs["detected_columns"] = {
        "BNU ID": bnu_column,
        "Name": name_column or full_name_column,
        "Surname": surname_column or "(not provided)",
        "Attendance": attendance_column,
        "Campus": campus_column or "(not provided)",
        "Group Ref": group_column or "(not provided)",
    }

    return prepared


def safe_filename_part(value):
    text = clean_text(value)
    text = re.sub(r'[<>:"/\\|?*]+', "", text)
    text = re.sub(r"\s+", "_", text)
    return text.strip("._") or "Unknown"


def fill_labeled_cell(doc, label_matcher, value):
    for table in doc.tables:
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                if label_matcher(cell.text):
                    if cell_index + 1 < len(row.cells):
                        row.cells[cell_index + 1].text = value
                        return True
                    if row_index + 1 < len(table.rows):
                        table.rows[row_index + 1].cells[cell_index].text = value
                        return True
    return False


def fill_student_details(doc, student_name, bnu_id, campus):
    filled_name = fill_labeled_cell(
        doc,
        lambda text: "name" in normalize_label(text) and "surname" not in normalize_label(text),
        student_name,
    )
    filled_bnu = fill_labeled_cell(doc, is_bnu_header, bnu_id)
    campus = clean_text(campus)
    filled_campus = fill_labeled_cell(doc, is_campus_header, campus) if campus else True

    if doc.tables:
        table = doc.tables[0]
        if len(table.rows) > 6 and len(table.rows[4].cells) > 1:
            if not filled_name:
                table.rows[4].cells[1].text = student_name
            if not filled_bnu:
                table.rows[5].cells[1].text = bnu_id
            if campus and not filled_campus:
                table.rows[6].cells[1].text = campus


def category_key(text):
    label = normalize_label(text)
    if "excellent" in label:
        return "Excellent attendance"
    if "very good" in label:
        return "Very good attendance"
    if "could be better" in label or "below" in label or "poor" in label:
        return "Attendance could be better"
    if re.search(r"\bgood\b", label):
        return "Good attendance"
    return None


def fill_attendance_category(doc, attendance_category):
    category_cells = {}

    for table in doc.tables:
        for row in table.rows:
            for cell_index, cell in enumerate(row.cells):
                row_category = category_key(cell.text)
                if not row_category or len(row.cells) <= 1:
                    continue

                target_index = cell_index + 1 if cell_index + 1 < len(row.cells) else cell_index - 1
                target_cell = row.cells[target_index]
                target_cell.text = ""
                category_cells[row_category] = target_cell

    if attendance_category in category_cells:
        category_cells[attendance_category].text = "Yes"
        return

    if len(doc.tables) > 1:
        attendance_table = doc.tables[1]
        row_map = {
            "Excellent attendance": 1,
            "Very good attendance": 2,
            "Good attendance": 3,
            "Attendance could be better": 4,
        }
        for row_number in row_map.values():
            if len(attendance_table.rows) > row_number and len(attendance_table.rows[row_number].cells) > 1:
                attendance_table.rows[row_number].cells[1].text = ""
        row_number = row_map[attendance_category]
        if len(attendance_table.rows) > row_number and len(attendance_table.rows[row_number].cells) > 1:
            attendance_table.rows[row_number].cells[1].text = "Yes"


def generate_reports(df, template_file, output_format, group_by, libreoffice_path):
    """Generate reports for all students"""
    zip_buffer = BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        
        for index, row in df.iterrows():
            student_name = clean_text(row["Student Name"])
            bnu_id = clean_bnu_id(row["BNU ID"])
            campus = clean_text(row["Campus"])
            attendance_category = row["Attendance Category"]

            if group_by:
                student_group = safe_filename_part(row["Group Ref"])
            
            # Create document from template
            doc = Document(template_file)
            fill_student_details(doc, student_name, bnu_id, campus)
            fill_attendance_category(doc, attendance_category)
            
            # Determine file path in zip
            surname = safe_filename_part(row["Surname"])
            first_name = safe_filename_part(row["Name"])
            if group_by:
                base_path = f"{student_group}/{safe_filename_part(bnu_id)}_{surname}_{first_name}_Attendance_Report"
            else:
                base_path = f"{safe_filename_part(bnu_id)}_{surname}_{first_name}_Attendance_Report"
            
            # Save as DOCX or convert to PDF
            if output_format == "DOCX":
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp:
                    doc.save(temp.name)
                    zip_file.write(temp.name, f"{base_path}.docx")
                    os.unlink(temp.name)
            
            else:  # PDF
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_docx:
                    doc.save(temp_docx.name)
                    
                    # Create temporary output directory
                    temp_dir = tempfile.mkdtemp()
                    
                    try:
                        # Convert to PDF using LibreOffice
                        command = [
                            libreoffice_path,
                            '--headless',
                            '--convert-to',
                            'pdf',
                            '--outdir',
                            temp_dir,
                            temp_docx.name
                        ]
                        
                        subprocess.run(
                            command,
                            capture_output=True,
                            text=True,
                            check=True,
                            timeout=30
                        )
                        
                        # Get PDF filename
                        pdf_name = os.path.splitext(os.path.basename(temp_docx.name))[0] + '.pdf'
                        pdf_path = os.path.join(temp_dir, pdf_name)
                        
                        # Add PDF to zip
                        if os.path.exists(pdf_path):
                            zip_file.write(pdf_path, f"{base_path}.pdf")
                            os.unlink(pdf_path)
                        
                    finally:
                        # Cleanup
                        os.unlink(temp_docx.name)
                        if os.path.exists(temp_dir):
                            shutil.rmtree(temp_dir)
    
    zip_buffer.seek(0)
    return zip_buffer

# Check LibreOffice availability
libreoffice_path = check_libreoffice()
has_libreoffice = libreoffice_path is not None

st.title("📊 Attendance Report Generator")
st.markdown("Generate professional attendance reports for your students")

# Show status
if has_libreoffice:
    st.success("✅ PDF conversion available!")
else:
    st.info("ℹ️ **Cloud Version:** DOCX only. For PDF support, add `packages.txt` file (see instructions in sidebar).")

st.divider()

st.header("📁 Step 1: Upload Files")

col1, col2 = st.columns(2)

with col1:
    excel_file = st.file_uploader(
        "Upload Student Attendance Excel",
        type=['xlsx', 'xls'],
        help="Upload your Excel file with student attendance data"
    )

with col2:
    template_file = st.file_uploader(
        "Upload Report Template (DOCX)",
        type=['docx'],
        help="Upload the Word document template for reports"
    )

st.divider()

st.header("⚙️ Step 2: Configure Options")

col3, col4 = st.columns(2)

with col3:
    if has_libreoffice:
        output_format = st.selectbox(
            "Output Format",
            ["PDF", "DOCX"],
            help="Choose whether to generate PDF or DOCX files"
        )
    else:
        output_format = "DOCX"
        st.info("📄 Output: DOCX (PDF requires LibreOffice)")

with col4:
    group_by = st.checkbox(
        "Organize by Groups",
        value=True,
        help="Organize reports into folders by student groups"
    )

st.divider()

st.header("🚀 Step 3: Generate Reports")

if st.button("Generate Reports", type="primary", use_container_width=True):
    
    if excel_file and template_file:
        
        if output_format == "PDF" and not has_libreoffice:
            st.error("❌ PDF conversion requires LibreOffice. Please use DOCX format or add packages.txt file.")
        else:
            try:
                with st.spinner("Reading attendance data..."):
                    df = read_attendance_data(excel_file)
                
                st.info(f"✅ Loaded {len(df)} students")
                st.caption(
                    f"Detected sheet: {df.attrs['source_sheet']} | "
                    f"Header row: {df.attrs['header_row']} | "
                    f"Columns: {df.attrs['detected_columns']}"
                )
                
                if group_by:
                    groups = df['Group Ref'].unique()
                    st.info(f"📂 Found {len(groups)} groups")
                
                with st.spinner(f"Generating {len(df)} {output_format} reports... This may take a few minutes."):
                    progress_bar = st.progress(0)
                    zip_buffer = generate_reports(df, template_file, output_format, group_by, libreoffice_path)
                    progress_bar.progress(100)
                
                st.success(f"🎉 Successfully generated {len(df)} reports!")
                
                st.download_button(
                    label=f"⬇️ Download All Reports ({output_format})",
                    data=zip_buffer,
                    file_name=f"attendance_reports_{output_format.lower()}.zip",
                    mime="application/zip",
                    type="primary",
                    use_container_width=True
                )
                
            except Exception as e:
                st.error(f"❌ Error: {str(e)}")
                st.exception(e)
    
    else:
        st.warning("⚠️ Please upload both files before generating reports")

with st.sidebar:
    st.header("📖 Instructions")
    
    st.markdown("""
    ### How to Use:
    
    1. **Upload Files**
       - Student attendance Excel file
       - Report template (DOCX)
    
    2. **Choose Options**
       - PDF or DOCX output
       - Organize by groups or not
    
    3. **Generate**
       - Click "Generate Reports"
       - Wait for processing
       - Download ZIP file
    
    
    ### Requirements:
    
    - The Excel sheet must include a BNU ID column, a student name column, and an attendance column.
    - Column names are detected case-insensitively, so `BNU ID`, `bnu id`, `LIVE`, `Attendance %`, and similar formats work.
    - Campus, surname, and group columns are optional.
    
    ### Support:
    
    Contact Ayan (ayan.achakzai@magnacartacollege.ac.uk) for help!
    """)
