"""
ATTENDANCE REPORT GENERATOR - DIRECT TO PDF (FULLY FIXED)
==========================================================
This script generates attendance reports and saves them directly as PDF.
Includes automatic column name cleaning to handle Excel quirks.

Author: Ayan
Date: January 2026
"""

import pandas as pd
from docx import Document
import os
import subprocess
import tempfile

# ============================================================================
# CONFIGURATION
# ============================================================================

INPUT_EXCEL = '/Users/ayanachakzai/Downloads/attendance report/student_attendance.xlsx'
TEMPLATE_DOCX = '/Users/ayanachakzai/Downloads/attendance report/template.docx'
OUTPUT_DIR = '/Users/ayanachakzai/Downloads/attendance report/attendance_reports_pdf'

# Path to LibreOffice
LIBREOFFICE_PATH = '/Applications/LibreOffice.app/Contents/MacOS/soffice'

# ============================================================================
# CHECK IF LIBREOFFICE IS INSTALLED
# ============================================================================

if not os.path.exists(LIBREOFFICE_PATH):
    print("❌ ERROR: LibreOffice not found!")
    print(f"   Expected location: {LIBREOFFICE_PATH}")
    print("\n📥 Please install LibreOffice:")
    print("   1. Download from: https://www.libreoffice.org/download/download/")
    print("   2. Install it like any other Mac app")
    print("   3. Run this script again")
    exit(1)

# ============================================================================
# CREATE OUTPUT FOLDER
# ============================================================================

os.makedirs(OUTPUT_DIR, exist_ok=True)

print("📖 Reading student attendance data...")

df = pd.read_excel(INPUT_EXCEL, skiprows=1)

# ============================================================================
# CLEAN COLUMN NAMES
# ============================================================================

df.columns = df.columns.str.strip()

print("✅ Column names cleaned")
print(f"   Available columns: {df.columns.tolist()}\n")

# ============================================================================
# CLEAN AND PREPARE DATA
# ============================================================================

df = df.dropna(subset=['BNU ID'])

# WHAT CHANGED HERE?
# - BEFORE: df['Name '] (with space, because column had space)
# - AFTER: df['Name'] (no space, because we stripped column names)
df['Name'] = df['Name'].str.strip()

df['Surname'] = df['Surname'].str.strip()

# WHAT CHANGED HERE?
# - BEFORE: df.rename(columns={'Name ': 'Name'}) - renamed column
# - AFTER: No rename needed! Column is already 'Name' after strip

print(f"✅ Found {len(df)} students to process\n")

# ============================================================================
# GET UNIQUE GROUPS
# ============================================================================

unique_groups = df['Group Ref'].unique()
unique_groups = sorted(unique_groups)

print(f"📊 Found {len(unique_groups)} unique groups:")
for group in unique_groups:
    student_count = (df['Group Ref'] == group).sum()
    print(f"   • {group}: {student_count} students")

print()

# ============================================================================
# PROCESS EACH STUDENT - GENERATE PDF DIRECTLY
# ============================================================================

for index, row in df.iterrows():
    
    student_name = f"{row['Name']} {row['Surname']}"
    
    bnu_id = str(int(row['BNU ID']))
    
    campus = row['Campus']
    
    attendance = row['LIVE']
    
    student_group = row['Group Ref']

    attendance_percent = attendance * 100
   
    if attendance_percent >= 80:
        attendance_category = "Excellent attendance"
    elif attendance_percent >= 70:
        attendance_category = "Very good attendance"
    elif attendance_percent >= 60:
        attendance_category = "Good attendance"
    else:
        attendance_category = "Attendance could be better"
    
    print(f"📝 Processing {index+1}/{len(df)}: {student_name} ({bnu_id}) - {attendance_percent:.1f}%")
   
    # ========================================================================
    # CREATE THE WORD DOCUMENT
    # ========================================================================
    
    doc = Document(TEMPLATE_DOCX)
    
    table = doc.tables[0]
    
    table.rows[4].cells[1].text = student_name
    
    table.rows[5].cells[1].text = bnu_id
    
    table.rows[6].cells[1].text = campus
  
    attendance_table = doc.tables[1]
    
    attendance_table.rows[1].cells[1].text = ""
    attendance_table.rows[2].cells[1].text = ""
    attendance_table.rows[3].cells[1].text = ""
    attendance_table.rows[4].cells[1].text = ""
    
    if attendance_category == "Excellent attendance":
        attendance_table.rows[1].cells[1].text = "Yes"
    elif attendance_category == "Very good attendance":
        attendance_table.rows[2].cells[1].text = "Yes"
    elif attendance_category == "Good attendance":
        attendance_table.rows[3].cells[1].text = "Yes"
    else:
        attendance_table.rows[4].cells[1].text = "Yes"
   
    # ========================================================================
    # CREATE GROUP FOLDER
    # ========================================================================
    
    group_folder = os.path.join(OUTPUT_DIR, student_group)
    os.makedirs(group_folder, exist_ok=True)
    
    # ========================================================================
    # SAVE AS PDF DIRECTLY
    # ========================================================================
    
    try:
        # Create temporary .docx file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_docx:
            temp_docx_path = temp_docx.name
            doc.save(temp_docx_path)
        
        # Create PDF filename
        pdf_filename = f"{bnu_id}_{row['Surname']}_{row['Name']}_Attendance_Report.pdf"
        pdf_path = os.path.join(group_folder, pdf_filename)
        
        # Convert to PDF using LibreOffice
        command = [
            LIBREOFFICE_PATH,
            '--headless',
            '--convert-to',
            'pdf',
            '--outdir',
            group_folder,
            temp_docx_path
        ]
        
        result = subprocess.run(
            command,
            capture_output=True,
            text=True,
            check=True
        )
        
        # Delete temporary .docx
        os.unlink(temp_docx_path)
        
        # Rename PDF to our desired filename
        temp_pdf_name = os.path.splitext(os.path.basename(temp_docx_path))[0] + '.pdf'
        temp_pdf_path = os.path.join(group_folder, temp_pdf_name)
        
        if os.path.exists(temp_pdf_path):
            os.rename(temp_pdf_path, pdf_path)
        
        print(f"   ✅ Saved as PDF: {pdf_filename}\n")
        
    except subprocess.CalledProcessError as e:
        print(f"   ❌ Error converting to PDF: {e}\n")
        if os.path.exists(temp_docx_path):
            os.unlink(temp_docx_path)
            
    except Exception as e:
        print(f"   ❌ Unexpected error: {e}\n")
        if 'temp_docx_path' in locals() and os.path.exists(temp_docx_path):
            os.unlink(temp_docx_path)


# ============================================================================
# COMPLETION MESSAGE
# ============================================================================

print(f"\n✅ DONE! Successfully generated {len(df)} PDF reports")
print(f"📁 PDFs organized by {len(unique_groups)} groups in: {OUTPUT_DIR}/")
print("\n📂 Folder structure:")
print(f"   {OUTPUT_DIR}/")
for group in unique_groups:
    student_count = (df['Group Ref'] == group).sum()
    print(f"   ├── {group}/ ({student_count} PDFs)")
print("\n🎉 All reports saved directly as PDF!")


# ============================================================================
# END OF SCRIPT
# ============================================================================

# WHAT WERE THE BUGS?
# ====================
# 
# BUG 1: Column had non-breaking space (\xa0)
# - Column was '\xa0Group Ref' not ' Group Ref'
# - Fixed by: df.columns = df.columns.str.strip()
# 
# BUG 2: After stripping column names, code still used old names
# - Code tried to access 'Name ' (with space)
# - But after strip, it's just 'Name' (no space)
# - Fixed by: Using df['Name'] instead of df['Name ']
# 
# BUG 3: Tried to rename column that doesn't exist anymore
# - Code tried: df.rename(columns={'Name ': 'Name'})
# - But 'Name ' doesn't exist after strip
# - Fixed by: Removing the rename line (not needed anymore)
# 
# KEY LESSON:
# ===========
# When you clean column names with .strip(), you must update
# ALL references to those columns in your code!
