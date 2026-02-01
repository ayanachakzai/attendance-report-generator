"""
ATTENDANCE REPORT GENERATOR
===========================
This script reads student attendance data from an Excel file and generates
individual Word document reports for each student with their attendance category.

Author: Ayan
Date: January 2026
"""

# ============================================================================
# SECTION 1: IMPORTING LIBRARIES (Tools we need)
# ============================================================================

import pandas as pd
from docx import Document
import os


# ============================================================================
# SECTION 2: CONFIGURATION (Settings/Variables we'll use)
# ============================================================================

# IMPORTANT: Use the EXACT filename as it appears in your Downloads folder
# Check if your file has spaces or underscores!
INPUT_EXCEL = '/Users/ayanachakzai/Downloads/attendance report/Students Attendance list Oct-25 intake updated till 25 Jan 26.xlsx'

TEMPLATE_DOCX = '/Users/ayanachakzai/Downloads/attendance report/ATTENDANCE_REPORT-SST.docx'

# Fixed the missing slash here!
OUTPUT_DIR = '/Users/ayanachakzai/Downloads/attendance report/attendance_reports'


# ============================================================================
# SECTION 3: CREATE OUTPUT FOLDER
# ============================================================================

os.makedirs(OUTPUT_DIR, exist_ok=True)


# ============================================================================
# SECTION 4: READ THE EXCEL FILE
# ============================================================================

print("📖 Reading student attendance data...")

df = pd.read_excel(INPUT_EXCEL, skiprows=1)


# ============================================================================
# SECTION 5: CLEAN THE DATA
# ============================================================================

df = df.dropna(subset=['BNU ID'])

df['Name '] = df['Name '].str.strip()

df['Surname'] = df['Surname'].str.strip()

df = df.rename(columns={'Name ': 'Name'})

print(f"✅ Found {len(df)} students to process\n")


# ============================================================================
# SECTION 6: PROCESS EACH STUDENT (Main Loop)
# ============================================================================

for index, row in df.iterrows():
    
    # ========================================================================
    # SECTION 6A: EXTRACT STUDENT DATA
    # ========================================================================
    
    student_name = f"{row['Name']} {row['Surname']}"
    
    bnu_id = str(int(row['BNU ID']))
    
    campus = row['Campus']
    
    attendance = row['LIVE']
    
    # ========================================================================
    # SECTION 6B: CALCULATE ATTENDANCE PERCENTAGE
    # ========================================================================
    
    attendance_percent = attendance * 100
    
    # ========================================================================
    # SECTION 6C: DETERMINE ATTENDANCE CATEGORY
    # ========================================================================
    
    if attendance_percent >= 80:
        attendance_category = "Excellent attendance"
    elif attendance_percent >= 70:
        attendance_category = "Very good attendance"
    elif attendance_percent >= 60:
        attendance_category = "Good attendance"
    else:
        attendance_category = "Attendance could be better"
    
    # ========================================================================
    # SECTION 6D: SHOW PROGRESS
    # ========================================================================
    
    print(f"📝 Processing {index+1}/{len(df)}: {student_name} ({bnu_id}) - {attendance_percent:.1f}%")
    
    # ========================================================================
    # SECTION 6E: LOAD AND FILL TEMPLATE
    # ========================================================================
    
    doc = Document(TEMPLATE_DOCX)
    
    table = doc.tables[0]
    
    table.rows[4].cells[1].text = student_name
    
    table.rows[5].cells[1].text = bnu_id
    
    table.rows[6].cells[1].text = campus
    
    # ========================================================================
    # SECTION 6F: FILL ATTENDANCE CHECKBOX
    # ========================================================================
    
    attendance_table = doc.tables[1]
    
    # Clear all checkboxes first (set to empty string)
    attendance_table.rows[1].cells[1].text = ""
    attendance_table.rows[2].cells[1].text = ""
    attendance_table.rows[3].cells[1].text = ""
    attendance_table.rows[4].cells[1].text = ""
    
    # Now mark the correct checkbox with "Yes"
    if attendance_category == "Excellent attendance":
        attendance_table.rows[1].cells[1].text = "Yes"
    elif attendance_category == "Very good attendance":
        attendance_table.rows[2].cells[1].text = "Yes"
    elif attendance_category == "Good attendance":
        attendance_table.rows[3].cells[1].text = "Yes"
    else:  # "Attendance could be better"
        attendance_table.rows[4].cells[1].text = "Yes"
    
    # ========================================================================
    # SECTION 6G: SAVE THE DOCUMENT
    # ========================================================================
    
    output_filename = f"{bnu_id}_{row['Surname']}_{row['Name']}_Attendance_Report.docx"
    
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    
    doc.save(output_path)


# ============================================================================
# SECTION 7: COMPLETION MESSAGE
# ============================================================================

print(f"\n✅ DONE! Successfully generated {len(df)} attendance reports")
print(f"📁 Reports saved in: {OUTPUT_DIR}/")
print("\n🎉 You can now find all reports in the 'attendance_reports' folder!")